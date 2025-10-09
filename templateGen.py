"""
templateGen.py
Gavin Schultz 2025
Generates legal documents with information based off of a given template.
"""

import docx

"""
Find and replace text in a document. This is primarily used to replace templates (i.e. 
%SURVEYOR%)
doc: Open docx file
find_text: Text to replace
replace_text: Text to replace with
returns: n/a
"""
def doc_find_and_replace(doc, find_text, replace_text):
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if find_text in run.text:
                run.text = run.text.replace(find_text, replace_text)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if find_text in run.text:
                            run.text = run.text.replace(find_text, replace_text)

"""
Fills out a surveyor's aff using a specified template. This handles the creation of the generated document. 
path: template document to use
surveyor: surveyor field in document
ftp: ftp field in document
file: file number in document
drawing: drawing field in document
legaldesc: legaldesc field in document
startdate: startdate field in document
enddate: enddate field in document
surv_city: surv_city field in document
returns: n/a
"""
def generate_surveyor_aff(path, surveyor, ftp, file, drawing, legaldesc, startdate, enddate, surv_city,out):
    doc = docx.Document(path)

    doc_find_and_replace(doc,r"%SURVEYOR%",surveyor)
    doc_find_and_replace(doc,r"%FTP%",ftp)
    doc_find_and_replace(doc,r"%FILE%",file)
    doc_find_and_replace(doc,r"%DRAWING%",drawing)
    doc_find_and_replace(doc,r"%LEGALDESC%",legaldesc)
    doc_find_and_replace(doc,r"%STARTDATE%",startdate)
    doc_find_and_replace(doc,r"%ENDDATE%",enddate)
    doc_find_and_replace(doc,r"%SURVEYORCITY%",surv_city)

    doc.save(out)

"""
Fills out a consent with seal using a specified template. This handles the creation of the generated document. 
path: template document to use
surveyor: surveyor field in document
corp: corporation field in document
filenumber: filenumber field in document
legal: legal description field in document
plantype: plan type field in document
returns: n/a
"""
def generate_consent_wseal(path, corp, plantype, surveyor, filenumber, legal,out):
    doc = docx.Document(path)

    doc_find_and_replace(doc,r"%SURVEYOR%",surveyor)
    doc_find_and_replace(doc,r"%CORPORATION%",corp)
    doc_find_and_replace(doc,r"%FILENUMBER%",filenumber)
    doc_find_and_replace(doc,r"%LEGAL%",legal)
    doc_find_and_replace(doc,r"%PLANTYPE%",plantype)

    doc.save(out)