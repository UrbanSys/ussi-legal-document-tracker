"""
Service for document generation from templates.
Wraps the existing templateGen.py logic for use in the FastAPI routes.
"""
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt


class DocumentGeneratorService:
    """Handles generation of legal documents from templates."""

    @staticmethod
    def doc_find_and_replace(doc: Document, find_text: str, replace_text: str) -> None:
        """
        Find and replace text in a Word document.
        
        Args:
            doc: Open DOCX document
            find_text: Text to find
            replace_text: Text to replace with
        """
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if find_text in run.text:
                                run.text = run.text.replace(find_text, replace_text)

    @staticmethod
    def generate_surveyor_aff(
        template_path: str,
        output_path: str,
        surveyor: str,
        ftp: str,
        file_num: str,
        drawing: str,
        legal_desc: str,
        start_date: str,
        end_date: str,
        surveyor_city: str,
    ) -> None:
        """
        Generate a surveyor's affidavit document from a template.
        
        Args:
            template_path: Path to the template DOCX file
            output_path: Path where the generated document will be saved
            surveyor: Surveyor name
            ftp: FTP number
            file_num: File number
            drawing: Drawing identifier
            legal_desc: Legal description
            start_date: Start date
            end_date: End date
            surveyor_city: City where surveyor is based
        """
        doc = Document(template_path)

        DocumentGeneratorService.doc_find_and_replace(doc, r"%SURVEYOR%", surveyor)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%FTP%", ftp)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%FILE%", file_num)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%DRAWING%", drawing)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%LEGALDESC%", legal_desc)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%STARTDATE%", start_date)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%ENDDATE%", end_date)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%SURVEYORCITY%", surveyor_city)

        doc.save(output_path)

    @staticmethod
    def generate_consent_with_seal(
        template_path: str,
        output_path: str,
        corporation: str,
        plan_type: str,
        surveyor: str,
        file_number: str,
        legal_desc: str,
    ) -> None:
        """
        Generate a consent with seal document from a template.
        
        Args:
            template_path: Path to the template DOCX file
            output_path: Path where the generated document will be saved
            corporation: Corporation/company name
            plan_type: Type of plan
            surveyor: Surveyor name
            file_number: File number
            legal_desc: Legal description
        """
        doc = Document(template_path)

        DocumentGeneratorService.doc_find_and_replace(doc, r"%SURVEYOR%", surveyor)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%CORPORATION%", corporation)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%FILENUMBER%", file_number)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%LEGAL%", legal_desc)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%PLANTYPE%", plan_type)

        doc.save(output_path)

    @staticmethod
    def generate_general_doc(
        template_path: str,
        output_path: str,
        corporation: str,
        plan_type: str,
        surveyor: str,
        file_number: str,
        legal_desc: str,
        doc_number: str,
    ) -> None:
        """
        Generate a general legal document from a template.
        
        Args:
            template_path: Path to the template DOCX file
            output_path: Path where the generated document will be saved
            corporation: Corporation/company name
            plan_type: Type of plan
            surveyor: Surveyor name
            file_number: File number
            legal_desc: Legal description
            doc_number: Document number
        """
        doc = Document(template_path)

        DocumentGeneratorService.doc_find_and_replace(doc, r"%SURVEYOR%", surveyor)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%CORPORATION%", corporation)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%FILENUMBER%", file_number)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%LEGAL%", legal_desc)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%PLANTYPE%", plan_type)
        DocumentGeneratorService.doc_find_and_replace(doc, r"%DOCNUMBER%", doc_number)

        doc.save(output_path)
